import re
import io
import pdfplumber
import docx
import spacy
from typing import Optional

# Load spaCy English model
# This gives us NLP capabilities — named entity recognition, etc.
try:
    nlp = spacy.load("en_core_web_lg")
except OSError:
    nlp = spacy.load("en_core_web_sm")


# ── SKILLS DATABASE ──
# This is our master list of skills to look for in resumes
SKILLS_DATABASE = [
    # Programming Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "php",
    "swift", "kotlin", "go", "rust", "scala", "r", "matlab", "perl",

    # Web Development
    "html", "css", "react", "angular", "vue", "nodejs", "express", "django",
    "flask", "fastapi", "spring", "asp.net", "jquery", "bootstrap", "tailwind",
    "nextjs", "nuxtjs", "graphql", "rest api", "restful",

    # Databases
    "sql", "mysql", "postgresql", "mongodb", "redis", "sqlite", "oracle",
    "cassandra", "elasticsearch", "dynamodb", "firebase",

    # Cloud & DevOps
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "jenkins",
    "git", "github", "gitlab", "ci/cd", "terraform", "ansible", "linux",

    # Data Science & AI
    "machine learning", "deep learning", "nlp", "natural language processing",
    "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
    "matplotlib", "seaborn", "power bi", "tableau", "excel", "spark",
    "hadoop", "data analysis", "data visualization", "statistics",

    # Mobile
    "android", "ios", "react native", "flutter", "xamarin",

    # Other Tech
    "agile", "scrum", "jira", "rest", "microservices", "blockchain",
    "cybersecurity", "networking", "linux", "bash", "shell scripting",

    # Soft Skills
    "communication", "leadership", "teamwork", "problem solving",
    "project management", "time management", "critical thinking",
]


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract all text from a PDF file"""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return text.strip()


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract all text from a DOCX file"""
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"DOCX extraction error: {e}")
    return text.strip()


def extract_email(text: str) -> Optional[str]:
    """Find email address in text"""
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None


def extract_phone(text: str) -> Optional[str]:
    """Find phone number in text"""
    patterns = [
        r'\+?1?\s*\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',  # US format
        r'\+?\d{1,3}[\s.-]?\d{3,5}[\s.-]?\d{4,6}',         # International
        r'\d{10}',                                            # Plain 10 digits
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text)
        if matches:
            return matches[0].strip()
    return None


def extract_name(text: str) -> Optional[str]:
    """
    Extract candidate name using spaCy NER (Named Entity Recognition).
    spaCy can identify PERSON entities in text.
    """
    # Process first 500 characters (name is usually at the top)
    doc = nlp(text[:500])

    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            # Filter out very short or very long names
            if 2 <= len(name.split()) <= 4:
                return name

    # Fallback: first line of resume often contains the name
    first_line = text.strip().split('\n')[0].strip()
    if len(first_line) < 50 and len(first_line) > 2:
        return first_line

    return None


def extract_skills(text: str) -> list[str]:
    """
    Find all skills mentioned in the resume by matching
    against our skills database.
    """
    text_lower = text.lower()
    found_skills = []

    for skill in SKILLS_DATABASE:
        # Use word boundary matching to avoid partial matches
        # e.g., "r" shouldn't match "react"
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            # Store with proper capitalization
            found_skills.append(skill.title())

    # Remove duplicates while preserving order
    seen = set()
    unique_skills = []
    for skill in found_skills:
        if skill.lower() not in seen:
            seen.add(skill.lower())
            unique_skills.append(skill)

    return unique_skills


def extract_experience_years(text: str) -> float:
    """
    Extract total years of experience from resume text.
    Looks for patterns like "5 years", "3+ years experience", etc.
    """
    text_lower = text.lower()
    patterns = [
        r'(\d+)\+?\s*years?\s*of\s*experience',
        r'(\d+)\+?\s*years?\s*experience',
        r'experience\s*of\s*(\d+)\+?\s*years?',
        r'(\d+)\+?\s*years?\s*in\s*the\s*industry',
        r'over\s*(\d+)\s*years?',
    ]

    years_found = []
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            try:
                years_found.append(float(match))
            except ValueError:
                pass

    if years_found:
        return max(years_found)  # Return highest number found

    # Alternative: count work experience entries by looking for year ranges
    year_pattern = r'(20\d{2}|19\d{2})\s*[-–]\s*(20\d{2}|19\d{2}|present|current)'
    year_matches = re.findall(year_pattern, text_lower)

    total_years = 0
    current_year = 2026
    for start, end in year_matches:
        try:
            start_yr = int(start)
            end_yr = current_year if end in ['present', 'current'] else int(end)
            total_years += max(0, end_yr - start_yr)
        except ValueError:
            pass

    return min(float(total_years), 40.0)  # Cap at 40 years


def extract_education(text: str) -> list[dict]:
    """Extract education details from resume"""
    education = []
    text_lower = text.lower()

    # Degree keywords to look for
    degrees = [
        "bachelor", "b.tech", "b.e", "b.sc", "b.com", "ba", "bs",
        "master", "m.tech", "m.e", "m.sc", "mba", "ms", "ma",
        "phd", "ph.d", "doctorate", "diploma", "associate",
        "10th", "12th", "ssc", "hsc", "intermediate"
    ]

    lines = text.split('\n')
    for i, line in enumerate(lines):
        line_lower = line.lower()
        for degree in degrees:
            if degree in line_lower:
                education.append({
                    "degree": line.strip(),
                    "details": lines[i+1].strip() if i+1 < len(lines) else ""
                })
                break

    return education[:5]  # Return max 5 education entries


def parse_resume(file_content: bytes, file_extension: str) -> dict:
    """
    Main function — parses a resume file and returns
    all extracted information as a dictionary.
    """
    # Step 1: Extract raw text based on file type
    if file_extension == ".pdf":
        raw_text = extract_text_from_pdf(file_content)
    elif file_extension == ".docx":
        raw_text = extract_text_from_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

    if not raw_text:
        return {
            "raw_text": "",
            "candidate_name": None,
            "candidate_email": None,
            "candidate_phone": None,
            "skills": [],
            "experience_years": 0.0,
            "education": [],
            "work_experience": [],
            "certifications": []
        }

    # Step 2: Extract all fields
    name = extract_name(raw_text)
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    skills = extract_skills(raw_text)
    experience_years = extract_experience_years(raw_text)
    education = extract_education(raw_text)

    print(f"✓ Parsed resume: {name} | Skills: {len(skills)} | Exp: {experience_years}yr")

    return {
        "raw_text": raw_text,
        "candidate_name": name,
        "candidate_email": email,
        "candidate_phone": phone,
        "skills": skills,
        "experience_years": experience_years,
        "education": education,
        "work_experience": [],
        "certifications": []
    }